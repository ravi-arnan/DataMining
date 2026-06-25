"""Dokumen Word: Ringkasan & penjelasan mudah Confusion Matrix.

Merangkum artikel Rina (Medium) "Memahami Confusion Matrix: Accuracy, Precision,
Recall, Specificity, dan F1-Score" menjadi catatan belajar yang mudah dipahami,
dengan menyematkan dua diagram dari artikel (tabel confusion matrix & ilustrasi
kucing/anjing).

Pakai:  .venv/bin/python docs/build_confusion.py
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, 'figures', 'cm')
OUT = os.path.join(HERE, 'Ringkasan_Confusion_Matrix.docx')

FONT = 'Times New Roman'
ACCENT = RGBColor(0x1F, 0x39, 0x64)
SRC = ('Sumber: Rina, "Memahami Confusion Matrix: Accuracy, Precision, Recall, '
       'Specificity, dan F1-Score untuk Evaluasi Model Klasifikasi", Medium. '
       'https://esairina.medium.com/memahami-confusion-matrix-accuracy-precision-'
       'recall-specificity-dan-f1-score-610d4f0db7cf')

doc = Document()
normal = doc.styles['Normal']
normal.font.name = FONT
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15


def sf(run, size=12, bold=False, italic=False, color=None, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.append(rf)
    rf.set(qn('w:eastAsia'), font)


def title(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run(t), size=17, bold=True, color=ACCENT)


def center(t, size=12, italic=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run(t), size=size, italic=italic)


def h1(t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12)
    sf(p.add_run(t), size=14, bold=True, color=ACCENT)


def h2(t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
    sf(p.add_run(t), size=12.5, bold=True)


def para(t):
    p = doc.add_paragraph(); sf(p.add_run(t))


def bullet(t):
    p = doc.add_paragraph(style='List Bullet'); sf(p.add_run(t))


def numbered(t):
    p = doc.add_paragraph(style='List Number'); sf(p.add_run(t))


def formula(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run(t), size=12, bold=True, italic=True, color=ACCENT)


def image(fname, width=5.0, caption=None):
    path = os.path.join(FIG, fname)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            center(caption, size=10, italic=True)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        sf(c.paragraphs[0].add_run(h), size=11, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            sf(cells[i].paragraphs[0].add_run(str(v)), size=11)
    doc.add_paragraph()


# ============================================================ COVER
title('Memahami Confusion Matrix')
center('Accuracy · Precision · Recall · Specificity · F1-Score', size=12)
center('Catatan Ringkas untuk Evaluasi Model Klasifikasi', size=12, italic=True)
center('Data Mining — Kelompok 10', size=11)
doc.add_paragraph()
para('Dokumen ini merangkum artikel sumber menjadi catatan belajar yang mudah '
     'dipahami: apa itu confusion matrix, empat komponennya, lima metrik '
     'evaluasi beserta rumus dan contoh perhitungannya, serta panduan memilih '
     'metrik yang tepat.')

# ============================================================ 1
h1('1. Kenapa Perlu Evaluasi Model?')
para('Evaluasi model menjawab pertanyaan: "Seberapa baik model kita dalam '
     'membuat prediksi?" Analoginya seperti latihan menendang bola — kita '
     'mencatat berapa kali bola masuk gawang (prediksi benar) dan berapa kali '
     'meleset (prediksi salah), lalu memakai catatan itu untuk memperbaiki '
     'kemampuan. Dalam data mining / machine learning, evaluasi memastikan '
     'model bekerja tepat, misalnya membedakan teks positif dan negatif.')
para('Untuk model klasifikasi, ukuran yang umum dipakai adalah accuracy, '
     'precision, recall, specificity, dan F1-score — semuanya dihitung dari '
     'sebuah tabel bernama confusion matrix.')

# ============================================================ 2
h1('2. Apa Itu Confusion Matrix?')
para('Confusion matrix (matriks kebingungan) adalah tabel yang membandingkan '
     'hasil PREDIKSI model dengan KENYATAAN sebenarnya pada data uji. Dari tabel '
     'ini kita bisa melihat di mana model benar dan di mana model keliru.')
image('img-002.png', width=4.2, caption='Struktur confusion matrix (TP, FP, FN, TN)')

h2('Empat Komponen Inti')
para('Bayangkan kasus deteksi: kelas "Positif" = yang ingin dideteksi, '
     '"Negatif" = sisanya.')
numbered('True Positive (TP): model menebak POSITIF dan kenyataannya memang '
         'POSITIF. (tebakan benar)')
numbered('True Negative (TN): model menebak NEGATIF dan kenyataannya memang '
         'NEGATIF. (tebakan benar)')
numbered('False Positive (FP): model menebak POSITIF padahal kenyataannya '
         'NEGATIF. (salah — "alarm palsu", Type I error)')
numbered('False Negative (FN): model menebak NEGATIF padahal kenyataannya '
         'POSITIF. (salah — "kelewatan", Type II error)')
para('Cara mudah mengingat: kata KEDUA (Positive/Negative) = apa yang DITEBAK '
     'model; kata PERTAMA (True/False) = apakah tebakan itu BENAR.')

h2('Contoh Visual: Kucing vs Anjing')
para('Kelas positif = KUCING, kelas negatif = ANJING. Model menebak 6 kucing '
     'benar (TP), 11 anjing benar (TN), 2 anjing yang dikira kucing (FP), dan 1 '
     'kucing yang dikira anjing (FN).')
image('img-003.png', width=5.2, caption='Ilustrasi confusion matrix kucing/anjing (dari artikel)')

# ============================================================ 3
h1('3. Contoh Kasus Utama: Review Film "Upin & Ipin"')
para('Artikel memakai 15 ulasan film dengan label sebenarnya (Positif/Negatif) '
     'dan hasil prediksi model. Setelah dibandingkan, diperoleh confusion '
     'matrix berikut (positif = ulasan memuji):')
table(['', 'Prediksi Positif', 'Prediksi Negatif'],
      [['Sebenarnya Positif', '6 (TP)', '3 (FN)'],
       ['Sebenarnya Negatif', '2 (FP)', '4 (TN)']])
para('Jadi: TP = 6, TN = 4, FP = 2, FN = 3, dengan total 15 data. '
     'Angka-angka inilah yang dipakai untuk semua metrik di bawah.')

# ============================================================ 4
h1('4. Lima Metrik Evaluasi')

h2('4.1 Accuracy (Akurasi)')
para('Mengukur seberapa banyak prediksi yang benar dari SELURUH data '
     '(positif maupun negatif).')
formula('Accuracy = (TP + TN) / (TP + TN + FP + FN)')
formula('= (6 + 4) / 15 = 10/15 = 0,67  (67%)')
para('Artinya 67% ulasan diprediksi benar. Catatan penting: akurasi bisa '
     'menyesatkan pada data TIDAK SEIMBANG (mis. 95% positif), karena model '
     'yang asal menebak kelas mayoritas pun terlihat "akurat".')

h2('4.2 Precision (Presisi)')
para('Dari semua yang DITEBAK POSITIF, berapa yang benar-benar positif? '
     'Fokus menekan kesalahan False Positive (alarm palsu).')
formula('Precision = TP / (TP + FP)')
formula('= 6 / (6 + 2) = 6/8 = 0,75  (75%)')
para('Penting saat FP mahal/berbahaya, mis. menandai email penting sebagai spam.')

h2('4.3 Recall / Sensitivitas')
para('Dari semua yang SEBENARNYA POSITIF, berapa yang berhasil ditemukan model? '
     'Fokus menekan kesalahan False Negative (yang kelewatan). Analogi: mencari '
     'jarum di tumpukan jerami — recall = seberapa banyak jarum yang ketemu.')
formula('Recall = TP / (TP + FN)')
formula('= 6 / (6 + 3) = 6/9 = 0,67  (67%)')
para('Penting saat FN berbahaya, mis. melewatkan pasien yang benar-benar sakit.')

h2('4.4 Specificity (Spesifisitas)')
para('Kebalikan recall: dari semua yang SEBENARNYA NEGATIF, berapa yang benar '
     'dikenali sebagai negatif.')
formula('Specificity = TN / (TN + FP)')
formula('= 4 / (4 + 2) = 4/6 = 0,67  (67%)')
para('Penting saat kita ingin memastikan kelas negatif tidak salah ditandai '
     'positif, mis. tes deteksi narkoba.')

h2('4.5 F1-Score')
para('Rata-rata harmonis dari Precision dan Recall — satu angka yang menyeimbangkan '
     'keduanya. Berguna saat FP dan FN sama-sama penting.')
formula('F1 = 2 × (Recall × Precision) / (Recall + Precision)')
formula('= 2 × (0,67 × 0,75) / (0,67 + 0,75) ≈ 0,71  (71%)')
para('F1 tinggi hanya jika precision DAN recall sama-sama cukup tinggi, sehingga '
     'lebih jujur daripada akurasi pada data tidak seimbang.')

h2('Ringkasan Hasil Contoh Upin & Ipin')
table(['Metrik', 'Rumus', 'Nilai'],
      [['Accuracy', '(TP+TN)/total', '0,67 (67%)'],
       ['Precision', 'TP/(TP+FP)', '0,75 (75%)'],
       ['Recall', 'TP/(TP+FN)', '0,67 (67%)'],
       ['Specificity', 'TN/(TN+FP)', '0,67 (67%)'],
       ['F1-Score', '2·(R·P)/(R+P)', '0,71 (71%)']])

# ============================================================ 5
h1('5. Panduan Memilih Metrik')
table(['Metrik', 'Pakai ketika...', 'Contoh'],
      [['Accuracy', 'kelas seimbang & semua kelas sama penting', 'prediksi cuaca: hujan / tidak'],
       ['Precision', 'ingin menekan False Positive', 'deteksi email spam'],
       ['Recall', 'ingin menekan False Negative', 'deteksi penyakit / kanker'],
       ['Specificity', 'ingin prediksi negatif benar-benar tepat', 'tes deteksi narkoba'],
       ['F1-Score', 'butuh keseimbangan Precision & Recall', 'deteksi berita clickbait']])
para('Tidak ada metrik yang "selalu terbaik". Seperti tukang memilih alat sesuai '
     'tujuan, metrik terbaik adalah yang paling relevan dengan tujuan model.')

# ============================================================ 6
h1('6. Kesimpulan')
numbered('Confusion matrix adalah dasar semua metrik: ia memetakan prediksi vs '
         'kenyataan ke dalam TP, TN, FP, FN.')
numbered('Accuracy mengukur kebenaran keseluruhan, tetapi bisa menipu pada data '
         'tidak seimbang.')
numbered('Precision menekan False Positive; Recall menekan False Negative — '
         'keduanya sering bertolak belakang.')
numbered('Specificity menilai ketepatan pada kelas negatif.')
numbered('F1-Score menyeimbangkan Precision dan Recall dalam satu angka.')
numbered('Pilih metrik berdasarkan TUJUAN dan jenis kesalahan yang paling '
         'merugikan dalam kasus kita.')

doc.add_paragraph()
p = doc.add_paragraph(); sf(p.add_run(SRC), size=9, italic=True)

doc.save(OUT)
print('Tersimpan:', OUT)
