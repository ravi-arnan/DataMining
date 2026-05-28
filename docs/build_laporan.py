"""Laporan UTS Data Mining Kelompok 10 - Apriori, K-Means, Agglomerative.
Format mengikuti template Kelompok 2 (UTS multi-metode dengan BAB III rumus +
BAB IV pembahasan per metode + BAB V penutup).

Pakai:  python build_laporan.py passA   (lalu konversi PDF)
        python build_laporan.py passB   (baca PDF passA, tulis docx final)
"""
import os
import sys
import subprocess
import numpy as np
import pandas as pd
from itertools import combinations
from scipy.spatial.distance import pdist, squareform
from scipy.cluster.hierarchy import linkage

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import (WD_ALIGN_PARAGRAPH, WD_LINE_SPACING,
                            WD_TAB_ALIGNMENT, WD_TAB_LEADER)
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, 'figures')
TMP_DOCX = os.path.join(HERE, '_passA.docx')
TMP_PDF = os.path.join(HERE, '_passA.pdf')
OUT_DOCX = os.path.join(HERE, 'Laporan_Kelompok10_UTS_DataMining.docx')

FONT = 'Times New Roman'
MONO = 'Courier New'

# ============================================================ DATA APRIORI (Tugas 1)
transaksi = {
    1:  {'susu', 'gula', 'teh'},
    2:  {'teh', 'gula', 'roti'},
    3:  {'teh', 'gula'},
    4:  {'susu', 'roti'},
    5:  {'susu', 'gula', 'roti'},
    6:  {'teh', 'gula'},
    7:  {'gula', 'kopi', 'susu'},
    8:  {'gula', 'kopi', 'susu'},
    9:  {'susu', 'roti', 'kopi'},
    10: {'gula', 'teh', 'kopi'},
}
produk = ['susu', 'gula', 'teh', 'roti', 'kopi']
TOTAL_TRX = len(transaksi)


def freq(itemset):
    s = set(itemset)
    return sum(1 for t in transaksi.values() if s.issubset(t))


def sup(itemset):
    return freq(itemset) / TOTAL_TRX


F1 = [(it, freq([it]), sup([it])) for it in produk]
F2 = [(' & '.join(c), freq(c), sup(c)) for c in combinations(produk, 2)]
F3 = [(' & '.join(c), freq(c), sup(c)) for c in combinations(produk, 3)]

# Aturan asosiasi
rules = []
for combo in list(combinations(produk, 2)) + list(combinations(produk, 3)):
    items = list(combo)
    for r in range(1, len(items)):
        for A in combinations(items, r):
            B = tuple(sorted(set(items) - set(A)))
            sup_AB = sup(items); sup_A = sup(A)
            conf = sup_AB / sup_A if sup_A > 0 else 0
            rules.append({
                'A': ', '.join(sorted(A)), 'B': ', '.join(B),
                'sup': sup_AB, 'conf': conf, 'final': sup_AB * conf,
            })
rules.sort(key=lambda r: r['final'], reverse=True)
TOP_RULES = rules[:10]

# ============================================================ DATA K-MEANS (Tugas 2)
cust = pd.DataFrame({
    'CustID': list(range(1, 11)),
    'Age':    [41, 47, 33, 29, 47, 40, 38, 42, 26, 47],
    'Income': [19, 100, 57, 19, 253, 81, 56, 64, 18, 115],
})
Xc = cust[['Age', 'Income']].values


def euclid(a, b):
    return float(np.sqrt(np.sum((np.array(a) - np.array(b)) ** 2)))


def kmeans_manual(X, init, max_iter=20):
    cent = np.array(init, dtype=float)
    hist = []
    for it in range(max_iter):
        labels = np.array([int(np.argmin([euclid(x, c) for c in cent])) for x in X])
        hist.append({'iter': it + 1, 'cent': cent.copy(), 'labels': labels.copy()})
        new_cent = np.array([X[labels == k].mean(axis=0) if (labels == k).any() else cent[k]
                             for k in range(len(cent))])
        if np.allclose(new_cent, cent):
            break
        cent = new_cent
    return cent, labels, hist


# K-Means K=2 manual
KM_INIT = [Xc[0], Xc[1]]
km_cent2, km_lab2, km_hist2 = kmeans_manual(Xc, KM_INIT)

# Jarak iterasi 1 (C1=P1, C2=P2)
km_iter1 = []
for i, x in enumerate(Xc):
    d1 = euclid(x, KM_INIT[0]); d2 = euclid(x, KM_INIT[1])
    km_iter1.append((i + 1, int(x[0]), int(x[1]), round(d1, 2), round(d2, 2), 1 if d1 < d2 else 2))

# SSE per K (data ter-standarisasi via sklearn dipakai untuk Elbow)
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
Xcs = StandardScaler().fit_transform(Xc)
KM_SSE = []
for k in range(1, 8):
    km = KMeans(n_clusters=k, n_init=10, random_state=42).fit(Xcs)
    KM_SSE.append((k, round(km.inertia_, 4)))

# ============================================================ DATA AGGLOMERATIVE (Tugas 3)
data = pd.DataFrame({
    'Mhs':   ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J'],
    'Tugas': [1, 2, 2, 5, 6, 5, 9, 9, 8, 10],
    'Ujian': [2, 2, 3, 5, 5, 6, 9, 8, 9, 10],
})
X = data[['Tugas', 'Ujian']].values
names = data['Mhs'].values
dist_mat = squareform(pdist(X))
Z = linkage(X, method='single', metric='euclidean')
gaps = np.diff(Z[:, 2])
i_max = int(np.argmax(gaps))
optimal_k = len(X) - (i_max + 1)
cut_height = (Z[i_max, 2] + Z[i_max + 1, 2]) / 2


def trace_single(X, names):
    clusters = [[n] for n in names]
    M = squareform(pdist(X)).astype(float)
    np.fill_diagonal(M, np.inf)
    rows = []
    while len(clusters) > 1:
        i, j = np.unravel_index(np.argmin(M), M.shape)
        if i > j: i, j = j, i
        a = '{' + ','.join(clusters[i]) + '}'
        b = '{' + ','.join(clusters[j]) + '}'
        d = M[i, j]
        new = np.minimum(M[i, :], M[j, :]); new[i] = np.inf
        M[i, :] = new; M[:, i] = new
        M = np.delete(M, j, 0); M = np.delete(M, j, 1)
        clusters[i] = clusters[i] + clusters[j]; del clusters[j]
        rows.append((len(rows) + 1, a, b, round(float(d), 3), len(clusters)))
    return rows


merge_rows = trace_single(X, names)

# ============================================================ STYLE HELPERS

def set_base_style(doc):
    st = doc.styles['Normal']
    st.font.name = FONT; st.font.size = Pt(12); st.font.color.rgb = RGBColor(0, 0, 0)
    rf = st.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(a), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE; pf.space_after = Pt(0)


def style_heading(doc, name, size):
    st = doc.styles[name]
    st.font.name = FONT; st.font.size = Pt(size); st.font.bold = True
    st.font.color.rgb = RGBColor(0, 0, 0)
    rf = st.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(a), FONT)
    st.paragraph_format.space_before = Pt(12); st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    st.paragraph_format.keep_with_next = True


def add_token(p, token):
    r = p.add_run(' ' + token)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(1); r.font.name = FONT


def body(doc, text, indent=True, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    p.add_run(text)
    return p


def numbered(doc, n, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.add_run(f'{n}.\t{text}')
    return p


def plain_center(doc, text, bold=False, size=12, sb=0, sa=0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    return p


def bab(doc, reg, roman, title):
    tok = f'ZT{len(reg):03d}'
    p1 = doc.add_heading(level=1); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run(f'BAB {roman}')
    p2 = doc.add_heading(level=1); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(title); add_token(p2, tok)
    reg.append(dict(cat='toc', level=0, label=f'BAB {roman} {title}', token=tok))


def h2num(doc, reg, num, title):
    tok = f'ZT{len(reg):03d}'
    p = doc.add_heading(level=2); p.add_run(f'{num}\t{title}'); add_token(p, tok)
    reg.append(dict(cat='toc', level=1, label=f'{num} {title}', token=tok))


def h3num(doc, reg, num, title):
    tok = f'ZT{len(reg):03d}'
    p = doc.add_heading(level=3); p.add_run(f'{num}\t{title}'); add_token(p, tok)
    reg.append(dict(cat='toc', level=2, label=f'{num} {title}', token=tok))


def special_heading(doc, reg, title):
    tok = f'ZT{len(reg):03d}'
    p = plain_center(doc, title, bold=True, sa=12); add_token(p, tok)
    reg.append(dict(cat='toc', level=0, label=title, token=tok))


def caption(doc, reg, cat, label):
    tok = f'Z{cat[0].upper()}{len(reg):03d}'
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(label); r.bold = True; r.font.size = Pt(11); r.font.name = FONT
    add_token(p, tok)
    reg.append(dict(cat=cat, level=0, label=label, token=tok))


def add_image(doc, fname, width_cm):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(os.path.join(FIG, fname), width=Cm(width_cm))


import json as _json
try:
    with open(os.path.join(FIG, 'eq_widths.json')) as _f:
        EQ_WIDTHS = _json.load(_f)
except FileNotFoundError:
    EQ_WIDTHS = {}


def add_equation(doc, fname, width_cm=None):
    """Embed rumus PNG centered, default pakai lebar natural dari eq_widths.json."""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    key = fname.replace('eq_', '').replace('.png', '')
    if width_cm is None:
        width_cm = EQ_WIDTHS.get(key, 10.0)
    p.add_run().add_picture(os.path.join(FIG, fname), width=Cm(width_cm))


def keterangan(doc, items):
    """Items: list of (sym, desc)."""
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.add_run('Keterangan:')
    for sym, desc in items:
        q = doc.add_paragraph()
        q.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        q.paragraph_format.left_indent = Cm(1.0)
        q.paragraph_format.first_line_indent = Cm(-1.0)
        q.paragraph_format.space_after = Pt(0)
        r = q.add_run(sym); r.italic = True
        q.add_run(f' : {desc}')


# tables
def set_cell_bg(cell, color):
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_table_borders(tbl):
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), '000000')
        borders.append(e)
    tbl._tbl.tblPr.append(borders)


def cell(c, text, bold=False, size=10, left=False):
    c.text = ''
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if left else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text)); r.bold = bold; r.font.name = FONT; r.font.size = Pt(size)


def make_table(doc, headers, rows, size=10):
    tbl = doc.add_table(rows=1, cols=len(headers)); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)
    for j, h in enumerate(headers):
        cell(tbl.rows[0].cells[j], h, bold=True, size=size); set_cell_bg(tbl.rows[0].cells[j], 'D9D9D9')
    for r in rows:
        cs = tbl.add_row().cells
        for j, v in enumerate(r):
            cell(cs[j], v, size=size)
    return tbl


def code_box(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)
    c = tbl.cell(0, 0); c.text = ''
    first = True
    for line in code_text.split('\n'):
        p = c.paragraphs[0] if first else c.add_paragraph(); first = False
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else ' '); r.font.name = MONO; r.font.size = Pt(8.5)


def page_footer(section, fmt, start):
    sectPr = section._sectPr
    pg = sectPr.find(qn('w:pgNumType'))
    if pg is None:
        pg = OxmlElement('w:pgNumType'); sectPr.append(pg)
    pg.set(qn('w:fmt'), fmt); pg.set(qn('w:start'), str(start))
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), ' PAGE ')
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'), FONT); rf.set(qn('w:hAnsi'), FONT)
    col = OxmlElement('w:color'); col.set(qn('w:val'), '000000')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '24')
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'none')
    for el in (rf, col, sz, u): rpr.append(el)
    t = OxmlElement('w:t'); t.text = '1'
    r.append(rpr); r.append(t); fld.append(r)
    p._p.append(fld)


def toc_line(doc, label, page, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.left_indent = Cm(0.0 if level == 0 else 0.75 * level)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(f'{label}\t{page}')
    if level == 0:
        r.bold = True


def detect_pages(pdf_path, reg):
    out = subprocess.run(['pdftotext', '-layout', pdf_path, '-'],
                         capture_output=True, text=True).stdout
    pages_text = out.split('\x0c')
    tokpage = {}
    for i, pg in enumerate(pages_text):
        for e in reg:
            if e['token'] in pg and e['token'] not in tokpage:
                tokpage[e['token']] = i
    first_tok = next(e['token'] for e in reg if e['cat'] == 'toc')
    base = tokpage.get(first_tok, 0)
    return {tok: (p - base + 1) for tok, p in tokpage.items()}


# ============================================================ BODY

APRIORI_CODE = '''from itertools import combinations
import pandas as pd

# 10 transaksi swalayan, 5 produk
transaksi = {
    1:  {'susu','gula','teh'},   2: {'teh','gula','roti'}, 3: {'teh','gula'},
    4:  {'susu','roti'},          5: {'susu','gula','roti'}, 6: {'teh','gula'},
    7:  {'gula','kopi','susu'},   8: {'gula','kopi','susu'},
    9:  {'susu','roti','kopi'},  10: {'gula','teh','kopi'},
}
produk = ['susu','gula','teh','roti','kopi']
N = len(transaksi)

def freq(itemset):
    s = set(itemset)
    return sum(1 for t in transaksi.values() if s.issubset(t))

def support(itemset):
    return freq(itemset) / N

# F1, F2, F3
F1 = [(it, support([it])) for it in produk]
F2 = [(c, support(c)) for c in combinations(produk, 2)]
F3 = [(c, support(c)) for c in combinations(produk, 3)]

# Aturan asosiasi A -> B
def buat_aturan(itemset):
    items = list(itemset); aturan = []
    for r in range(1, len(items)):
        for A in combinations(items, r):
            B = tuple(sorted(set(items) - set(A)))
            sAB = support(items); sA = support(A)
            conf = sAB / sA if sA > 0 else 0
            aturan.append((A, B, sAB, conf, sAB * conf))
    return aturan

rules = []
for combo in list(combinations(produk, 2)) + list(combinations(produk, 3)):
    rules += buat_aturan(combo)
rules.sort(key=lambda r: r[4], reverse=True)
print('Top 5 aturan:', rules[:5])'''

KMEANS_CODE = '''import numpy as np
import pandas as pd
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler

data = pd.DataFrame({
    'Age':    [41, 47, 33, 29, 47, 40, 38, 42, 26, 47],
    'Income': [19, 100, 57, 19, 253, 81, 56, 64, 18, 115],
})
X = data[['Age', 'Income']].values

# K-Means manual K=2 dengan centroid awal P1 dan P2
def euclid(a, b):
    return np.sqrt(np.sum((a - b) ** 2))

def kmeans_manual(X, init, max_iter=20):
    cent = np.array(init, dtype=float)
    for it in range(max_iter):
        labels = np.array([np.argmin([euclid(x, c) for c in cent]) for x in X])
        new = np.array([X[labels == k].mean(axis=0) for k in range(len(cent))])
        if np.allclose(new, cent): break
        cent = new
    return cent, labels

cent2, lab2 = kmeans_manual(X, [X[0], X[1]])
print('Centroid K=2 (asli):', cent2)
print('Label:', lab2)

# K=3..6 + Elbow (data ter-standarisasi agar adil)
Xs = StandardScaler().fit_transform(X)
for k in range(1, 8):
    km = KMeans(n_clusters=k, n_init=10, random_state=42).fit(Xs)
    print(f'K={k}  SSE={km.inertia_:.4f}')'''

AGG_CODE = '''import numpy as np, pandas as pd
from scipy.spatial.distance import pdist, squareform
from scipy.cluster.hierarchy import linkage, fcluster

data = pd.DataFrame({
    'Mhs':   ['A','B','C','D','E','F','G','H','I','J'],
    'Tugas': [1, 2, 2, 5, 6, 5, 9, 9, 8, 10],
    'Ujian': [2, 2, 3, 5, 5, 6, 9, 8, 9, 10],
})
X = data[['Tugas','Ujian']].values
names = data['Mhs'].values

def agglomerative_single(X, names):
    clusters = [[n] for n in names]
    M = squareform(pdist(X)).astype(float); np.fill_diagonal(M, np.inf)
    while len(clusters) > 1:
        i, j = np.unravel_index(np.argmin(M), M.shape)
        if i > j: i, j = j, i
        print(f'Gabung {clusters[i]} + {clusters[j]}  jarak={M[i,j]:.3f}')
        new = np.minimum(M[i,:], M[j,:]); new[i] = np.inf  # rumus MIN
        M[i,:] = new; M[:,i] = new
        M = np.delete(M, j, 0); M = np.delete(M, j, 1)
        clusters[i] += clusters[j]; del clusters[j]

agglomerative_single(X, names)

# Verifikasi + jumlah klaster dari gap terbesar
Z = linkage(X, method='single')
gaps = np.diff(Z[:, 2]); i = int(np.argmax(gaps))
k = len(X) - (i + 1)
print('Jumlah klaster optimal =', k)
print('Label klaster:', fcluster(Z, t=k, criterion='maxclust'))'''


def build_body(doc, reg):
    # ============================================================ BAB I
    bab(doc, reg, 'I', 'PENDAHULUAN')
    body(doc, 'Bab I Pendahuluan ini membahas latar belakang, rumusan masalah, tujuan, manfaat, dan batasan masalah terkait implementasi tiga algoritma data mining yang dipelajari oleh kelompok, yaitu algoritma Apriori untuk penemuan aturan asosiasi pada data transaksi, algoritma K-Means untuk pengelompokan data pelanggan berdasarkan usia dan pendapatan, serta algoritma Agglomerative Hierarchical Clustering untuk pengelompokan data mahasiswa berdasarkan nilai akademik.')

    h2num(doc, reg, '1.1', 'Latar Belakang')
    body(doc, 'Perkembangan teknologi informasi menyebabkan jumlah data yang dihasilkan di berbagai bidang seperti pendidikan, kesehatan, bisnis, dan industri terus meningkat. Data dalam jumlah besar tersebut dapat menjadi sumber informasi penting apabila diolah dan dianalisis dengan tepat. Namun, banyaknya data yang tersedia sering sulit dianalisis secara manual sehingga diperlukan metode yang mampu menggali pola dan pengetahuan dari data tersebut. Salah satu bidang ilmu yang berfokus pada permasalahan ini adalah data mining.')
    body(doc, 'Data mining merupakan proses menemukan pola, hubungan, dan informasi penting dari kumpulan data berukuran besar menggunakan metode statistik, machine learning, dan basis data (Larose dan Larose, 2014). Dalam praktiknya, data mining memiliki beberapa teknik utama, antara lain asosiasi, klasifikasi, clustering, estimasi, dan forecasting. Pada laporan ini dikaji tiga algoritma fundamental yang merepresentasikan teknik-teknik tersebut beserta simulasi perhitungannya pada permasalahan yang spesifik.')
    body(doc, 'Pertama, algoritma Apriori digunakan untuk analisis asosiasi guna mengungkap keterkaitan kombinasi produk pada data transaksi pelanggan. Kedua, algoritma K-Means diterapkan sebagai teknik clustering untuk mengidentifikasi pengelompokan pelanggan berdasarkan tingkat usia (Age) dan pendapatan (Income). Ketiga, algoritma Agglomerative Hierarchical Clustering digunakan sebagai metode clustering hirarki untuk mengelompokkan mahasiswa berdasarkan Nilai Tugas dan Nilai Ujian, sekaligus menampilkan hirarki klaster dalam bentuk dendrogram. Melalui penerapan ketiga algoritma ini, diharapkan dapat dipahami konsep matematis, langkah komputasi, dan karakteristik hasil dari masing-masing metode secara komprehensif.')

    h2num(doc, reg, '1.2', 'Rumusan Masalah')
    body(doc, 'Berdasarkan latar belakang di atas, rumusan masalah dalam penulisan laporan ini adalah sebagai berikut.', indent=False)
    for i, t in enumerate([
        'Bagaimana proses pembentukan aturan asosiasi (2-itemset dan 3-itemset) pada data transaksi pelanggan menggunakan algoritma Apriori, beserta perhitungan support, confidence, dan nilai final aturan.',
        'Bagaimana cara menentukan jumlah klaster optimal dan proses pengelompokan data Age serta Income menggunakan algoritma K-Means dengan bantuan Elbow Method.',
        'Bagaimana langkah komputasi algoritma Agglomerative Hierarchical Clustering dalam membentuk hirarki klaster, serta bagaimana menentukan jumlah klaster final dari dendrogram.'], 1):
        numbered(doc, i, t)

    h2num(doc, reg, '1.3', 'Tujuan Penulisan')
    body(doc, 'Tujuan dari penulisan laporan eksperimen data mining ini adalah sebagai berikut.', indent=False)
    for i, t in enumerate([
        'Mengimplementasikan algoritma Apriori untuk membedah dan menghitung nilai support serta confidence keterkaitan antarbarang pada data transaksi pelanggan.',
        'Menerapkan metode clustering dengan algoritma K-Means untuk mengekstraksi segmentasi kelompok data Age dan Income berdasarkan kedekatan karakteristiknya.',
        'Membangun dan memahami tahapan komputasi algoritma Agglomerative Clustering, mulai dari perhitungan matriks jarak, penggabungan klaster, hingga penentuan jumlah klaster dari dendrogram.'], 1):
        numbered(doc, i, t)

    h2num(doc, reg, '1.4', 'Manfaat Penulisan')
    numbered(doc, 1, 'Bagi Penulis/Mahasiswa: memberikan fondasi teoritis dan teknis yang kuat, baik secara matematis maupun pemrograman (menggunakan Python, pustaka pandas, scipy, dan scikit-learn), terkait implementasi algoritma penambangan data mulai dari pencarian pola keranjang belanja (Apriori), pengelompokan data tak berlabel (K-Means dan Agglomerative).')
    numbered(doc, 2, 'Bagi Akademisi: dapat digunakan sebagai referensi terapan implementasi algoritma data mining yang memadukan teori probabilitas bersyarat, jarak Euclidean, dan teknik unsupervised learning.')
    numbered(doc, 3, 'Bagi Praktisi: menyediakan wawasan analitik dan prototipe dasar solusi pengambilan keputusan berbasis data, misalnya untuk merancang strategi cross-selling melalui market basket analysis, menentukan target segmentasi pelanggan melalui K-Means, hingga melakukan analisis hirarkis pengelompokan data observasi.')

    h2num(doc, reg, '1.5', 'Batasan Masalah')
    body(doc, 'Agar pembahasan lebih terfokus, penulisan laporan ini memiliki batasan masalah sebagai berikut.', indent=False)
    for i, t in enumerate([
        'Implementasi algoritma Apriori dibatasi pada pencarian kombinasi maksimal himpunan tiga barang (F3) menggunakan 10 sampel data transaksi dengan lima produk (susu, gula, teh, roti, kopi).',
        'Percobaan algoritma K-Means diterapkan pada 10 sampel data pelanggan dengan dua variabel (Age dan Income). Perhitungan manual dilakukan untuk K=2, sedangkan K=3 hingga K=6 dihitung dengan program serta dianalisis menggunakan Elbow Method.',
        'Implementasi algoritma Agglomerative Clustering dibatasi pada 10 data mahasiswa dengan dua atribut (Nilai Tugas dan Nilai Ujian), menggunakan jarak Euclidean dan metode single linkage. Penentuan jumlah klaster dilakukan dengan memotong dendrogram pada lompatan jarak terbesar.'], 1):
        numbered(doc, i, t)

    # ============================================================ BAB II
    doc.add_page_break()
    bab(doc, reg, 'II', 'TINJAUAN PUSTAKA')
    body(doc, 'Bab II Tinjauan Pustaka ini memaparkan landasan teori yang mendasari algoritma dan metode penambangan data (data mining) yang digunakan di dalam laporan. Topik yang dibahas mencakup pemahaman dasar mengenai data mining beserta teknik-tekniknya, konsep aturan asosiasi melalui algoritma Apriori, metode pengelompokan data tanpa label (clustering) menggunakan algoritma K-Means, hingga metode hierarchical clustering menggunakan algoritma Agglomerative.')

    h2num(doc, reg, '2.1', 'Data Mining')
    body(doc, 'Data mining adalah proses komputasi interdisipliner yang erat kaitannya dengan Knowledge Discovery in Database (KDD), yaitu proses penggalian pola tersembunyi dari kumpulan data berukuran besar untuk diubah menjadi pengetahuan yang berguna dalam pengambilan keputusan. Data mining bekerja dengan memadukan tiga fondasi keilmuan utama, yaitu statistik, kecerdasan buatan (machine learning), dan teknologi basis data. Beberapa teknik utama dalam data mining antara lain estimasi, prediksi, klasifikasi, clustering, dan asosiasi (Larose dan Larose, 2014).')

    h2num(doc, reg, '2.2', 'Association')
    body(doc, 'Aturan asosiasi (Association Rule) adalah salah satu metode utama dalam data mining yang digunakan untuk mencari dan menemukan pola kombinasi item (itemset) yang bermakna dan paling sering muncul secara bersamaan di dalam sekumpulan transaksi atau basis data yang besar. Teknik ini sering juga disebut sebagai affinity analysis atau market basket analysis. Dalam proses analitiknya, pencarian asosiasi selalu bergantung pada dua tolok ukur utama, yaitu nilai support dan nilai confidence. Support mengukur persentase seberapa sering suatu kombinasi item muncul di dalam seluruh transaksi, sedangkan confidence mengukur seberapa kuat hubungan keterkaitan antara dua item pada kondisi tertentu.')

    h2num(doc, reg, '2.3', 'Classification')
    body(doc, 'Klasifikasi dalam data mining adalah sebuah proses untuk menemukan model prediktif yang membagi sekumpulan data ke dalam kelompok atau kelas tertentu berdasarkan batasan yang telah ditentukan sebelumnya. Berbeda dengan estimasi yang memprediksi nilai numerik, klasifikasi beroperasi pada target variabel yang bersifat kategori. Beberapa algoritma klasifikasi yang dikenal antara lain decision tree (ID3, C4.5, CART), Naive Bayes, k-Nearest Neighbor, dan Support Vector Machine. Pada laporan ini, klasifikasi tidak dibahas lebih jauh karena fokus utama berada pada teknik asosiasi dan clustering.')

    h2num(doc, reg, '2.4', 'Clustering')
    body(doc, 'Klastering (clustering) adalah sebuah teknik analisis data di dalam data mining dan cabang utama dari pembelajaran mesin tanpa pengawasan (unsupervised learning) yang berfungsi untuk mengelompokkan objek-objek data yang tidak memiliki label (unlabeled data). Tujuan utama dari metode ini adalah mengungkapkan struktur bawaan atau kelas-kelas tersembunyi di dalam suatu himpunan data, di mana hubungan antar objeknya belum diketahui secara pasti sebelumnya. Dalam praktiknya, algoritma clustering bekerja dengan cara membagi kumpulan data ke dalam kelompok-kelompok terpisah sehingga objek dalam satu kelompok memiliki kesamaan karakteristik yang tinggi, sementara objek di kelompok lain memiliki sifat yang berbeda. Pendekatan clustering yang umum digunakan adalah partitional clustering seperti K-Means dan hierarchical clustering seperti Agglomerative Clustering.')

    h2num(doc, reg, '2.5', 'Estimation')
    body(doc, 'Estimasi dalam data mining adalah sebuah proses analitik prediktif yang berfungsi untuk menghitung, memperkirakan, atau menentukan nilai numerik secara presisi dari suatu variabel target berdasarkan pola hubungannya dengan satu atau lebih variabel prediktor. Berbeda dengan tugas klasifikasi yang memprediksi label kategori, proses estimasi murni berfokus pada prediksi nilai kontinu. Dalam penerapannya, metodologi estimasi membentang dari pendekatan statistik tradisional seperti regresi linier berganda hingga implementasi algoritma machine learning seperti Artificial Neural Network (ANN) dan ensemble method (Random Forest, XGBoost).')

    h2num(doc, reg, '2.6', 'Apriori')
    body(doc, 'Algoritma Apriori adalah salah satu algoritma dasar dalam data mining yang diusulkan oleh Agrawal dan Srikant pada tahun 1994, yang termasuk ke dalam jenis Association Rule Mining. Algoritma ini berfungsi untuk menemukan aturan asosiatif atau mengungkapkan keterkaitan antar kombinasi item dalam suatu basis data transaksional, yang mana aturan tersebut sering juga dikenal dengan sebutan Affinity Analysis atau Market Basket Analysis. Dalam penerapannya, algoritma ini bekerja dengan berfokus pada analisis pola frekuensi tinggi (frequent pattern mining) melalui pencarian dan pengkombinasian kumpulan item yang paling sering muncul secara bersamaan (frequent itemsets). Proses pencarian ini sangat bergantung pada perhitungan dua tolok ukur utama, yaitu nilai support (persentase kemunculan suatu kombinasi item di dalam keseluruhan transaksi) dan nilai confidence (tingkat kepastian atau kuatnya hubungan antar item dalam aturan tersebut).')

    h2num(doc, reg, '2.7', 'K-Means')
    body(doc, 'K-Means merupakan salah satu algoritma clustering yang paling banyak digunakan dalam analisis data tak terstruktur (unsupervised learning). Algoritma ini bekerja dengan cara mengelompokkan sekumpulan data ke dalam K klaster berdasarkan kemiripan fitur, di mana setiap data point akan dimasukkan ke dalam klaster dengan centroid terdekat menggunakan ukuran jarak, umumnya Euclidean distance. Proses ini pertama kali diperkenalkan oleh MacQueen pada 1967 dan kemudian dikembangkan lebih lanjut oleh Hartigan dan Wong pada 1979. K-Means tergolong dalam metode partitional clustering yang bersifat iteratif, di mana kualitas klaster diukur menggunakan fungsi objektif berupa minimisasi within-cluster sum of squares (WCSS) atau yang juga dikenal sebagai inertia. Penentuan jumlah klaster optimal (K) umumnya dilakukan menggunakan Elbow Method.')

    h2num(doc, reg, '2.8', 'Agglomerative Clustering')
    body(doc, 'Agglomerative Hierarchical Clustering merupakan salah satu metode hierarchical clustering dengan pendekatan bottom-up. Algoritma ini bekerja dengan memperlakukan setiap data sebagai satu klaster sendiri pada tahap awal, lalu pada setiap iterasi menggabungkan dua klaster yang memiliki jarak terdekat hingga seluruh data tergabung dalam satu klaster besar. Seluruh proses penggabungan dicatat sebagai hirarki klaster dan divisualisasikan dalam bentuk dendrogram. Berbeda dengan K-Means yang memerlukan penentuan K di awal, Agglomerative Clustering tidak memerlukan jumlah klaster yang ditetapkan di awal. Jumlah klaster final ditentukan dengan memotong dendrogram pada ketinggian tertentu, salah satunya melalui analisis lompatan (gap) jarak penggabungan terbesar. Metode jarak antar klaster (linkage) yang umum digunakan antara lain single linkage (minimum), complete linkage (maksimum), average linkage (rata-rata), dan ward linkage (minimalisasi kenaikan SSE).')

    # ============================================================ BAB III
    doc.add_page_break()
    bab(doc, reg, 'III', 'METODE PENELITIAN')
    body(doc, 'Bab III Metode Penelitian ini menguraikan langkah-langkah logis dan rancangan komputasi tiap algoritma untuk memecahkan setiap rumusan masalah. Pembahasan mencakup tahapan penemuan aturan asosiasi menggunakan algoritma Apriori (melalui ekstraksi item dan kalkulasi probabilitas support serta confidence), proses pengelompokan data dengan algoritma K-Means (melalui penentuan klaster optimal via metode Elbow dan iterasi centroid), serta proses pengelompokan hirarkis menggunakan algoritma Agglomerative Clustering (melalui perhitungan matriks jarak, penggabungan single linkage, dan analisis dendrogram).')

    # ----- 3.1 Apriori -----
    h2num(doc, reg, '3.1', 'Algoritma Apriori')
    body(doc, 'Metode pembentukan aturan asosiasi diawali dengan mengekstrak seluruh data keranjang transaksi untuk mengidentifikasi item unik tanpa duplikasi, sekaligus menghitung total populasi transaksi sebagai basis pembagi nilai. Tahap pembentukan kandidat dilakukan dengan menyusun seluruh kemungkinan kombinasi himpunan berukuran tertentu (seperti kombinasi dua atau tiga item) secara kombinatorial. Untuk mengetahui tingkat kepopuleran setiap kombinasi, dilakukan perhitungan frekuensi kemunculannya di dalam keseluruhan data transaksi yang kemudian dikonversi menjadi persentase nilai penunjang (support).')

    h3num(doc, reg, '3.1.1', 'Perhitungan Support')
    body(doc, 'Support adalah persentase kemunculan suatu itemset di dalam keseluruhan transaksi. Nilai ini dihitung dengan membagi frekuensi kemunculan itemset terhadap total jumlah transaksi (N).', indent=False)
    add_equation(doc, 'eq_apr_support.png')
    keterangan(doc, [
        ('Support(X)', 'persentase kemunculan itemset X dalam seluruh transaksi'),
        ('freq(X)', 'jumlah transaksi yang memuat seluruh item di X'),
        ('N', 'total jumlah transaksi'),
    ])

    h3num(doc, reg, '3.1.2', 'Perhitungan Confidence')
    body(doc, 'Confidence mengukur tingkat kepastian aturan A → B, yaitu seberapa sering item B muncul ketika item A muncul. Nilai confidence diperoleh dengan membagi support gabungan A dan B terhadap support A.', indent=False)
    add_equation(doc, 'eq_apr_confidence.png')
    keterangan(doc, [
        ('Confidence(A → B)', 'tingkat kepastian aturan A → B'),
        ('Support(A ∪ B)', 'support gabungan A dan B'),
        ('Support(A)', 'support sisi kiri aturan'),
    ])

    h3num(doc, reg, '3.1.3', 'Nilai Final Aturan')
    body(doc, 'Untuk meranking seluruh aturan, dihitung nilai final sebagai hasil kali support dan confidence. Aturan dengan nilai final tertinggi merupakan aturan yang paling kuat.', indent=False)
    add_equation(doc, 'eq_apr_final.png')
    keterangan(doc, [
        ('Final(A → B)', 'metrik akhir untuk meranking aturan A → B'),
    ])

    # ----- 3.2 K-Means -----
    h2num(doc, reg, '3.2', 'Algoritma K-Means')
    body(doc, 'K-Means diimplementasikan melalui beberapa tahapan sistematis. Pertama, ditentukan nilai k (jumlah klaster) yang optimal menggunakan metode Elbow, yaitu dengan memplot nilai SSE/WCSS terhadap berbagai nilai k dan memilih titik di mana penurunan SSE mulai melandai. Selanjutnya, algoritma diinisialisasi dengan menempatkan k centroid pada ruang fitur, kemudian setiap data point dihitung jaraknya terhadap seluruh centroid dan ditetapkan ke klaster dengan centroid terdekat. Setelah seluruh data dikelompokkan, posisi centroid diperbarui dengan menghitung rata-rata koordinat anggota klaster. Proses ini diulang secara iteratif hingga posisi centroid tidak lagi berubah secara signifikan.')

    h3num(doc, reg, '3.2.1', 'Penghitungan Jarak (Euclidean Distance)')
    body(doc, 'Jarak antara titik data dan centroid dihitung menggunakan rumus Euclidean distance pada ruang fitur berdimensi m.', indent=False)
    add_equation(doc, 'eq_km_euclidean.png')
    keterangan(doc, [
        ('d(x_i, c_k)', 'jarak antara data x_i dan centroid c_k'),
        ('x_ij', 'nilai atribut ke-j dari data x_i'),
        ('c_kj', 'nilai atribut ke-j dari centroid c_k'),
        ('m', 'jumlah atribut/fitur'),
    ])

    h3num(doc, reg, '3.2.2', 'Penetapan Klaster')
    body(doc, 'Setiap data point ditetapkan ke klaster dengan centroid terdekat menggunakan operasi argmin.', indent=False)
    add_equation(doc, 'eq_km_assign.png')
    keterangan(doc, [
        ('C(x_i)', 'label klaster yang ditetapkan untuk data x_i'),
        ('K', 'jumlah klaster'),
    ])

    h3num(doc, reg, '3.2.3', 'Pembaruan Centroid')
    body(doc, 'Setelah seluruh data dilabeli, posisi centroid setiap klaster diperbarui dengan menghitung rata-rata seluruh anggotanya.', indent=False)
    add_equation(doc, 'eq_km_centroid.png')
    keterangan(doc, [
        ('c_k', 'centroid klaster ke-k yang baru'),
        ('|C_k|', 'jumlah anggota klaster ke-k'),
    ])

    h3num(doc, reg, '3.2.4', 'Evaluasi dan Penentuan K Optimal (Elbow Method)')
    body(doc, 'Kualitas hasil clustering dievaluasi menggunakan total jarak kuadrat dalam klaster atau Sum of Squared Errors (SSE).', indent=False)
    add_equation(doc, 'eq_km_sse.png')
    keterangan(doc, [
        ('SSE', 'total jarak kuadrat anggota klaster terhadap centroidnya'),
        ('||x - c_k||', 'norma Euclidean antara data x dan centroid c_k'),
    ])
    body(doc, 'Elbow Method menghitung SSE untuk berbagai nilai K dan memilih nilai K pada titik di mana kurva SSE mulai melandai, yang menandakan penambahan K tidak lagi memberikan penurunan SSE yang signifikan.')

    # ----- 3.3 Agglomerative -----
    h2num(doc, reg, '3.3', 'Algoritma Agglomerative Clustering')
    body(doc, 'Agglomerative Clustering diimplementasikan melalui serangkaian tahapan komputasi yang sistematis dan berurutan. Pada tahap awal, setiap data diperlakukan sebagai satu klaster sendiri sehingga jumlah klaster awal sama dengan jumlah data (n). Selanjutnya, dihitung matriks jarak antar seluruh klaster dengan ukuran Euclidean distance. Pada setiap iterasi, dua klaster dengan jarak terdekat digabung, kemudian matriks jarak diperbarui menggunakan aturan linkage. Proses ini diulang hingga seluruh data tergabung menjadi satu klaster. Jumlah klaster final ditentukan dengan memotong dendrogram pada lompatan jarak penggabungan terbesar.')

    h3num(doc, reg, '3.3.1', 'Penghitungan Jarak Euclidean')
    body(doc, 'Jarak antara dua data dihitung menggunakan rumus jarak Euclidean.', indent=False)
    add_equation(doc, 'eq_ag_euclidean.png')
    keterangan(doc, [
        ('d(p, q)', 'jarak Euclidean antara titik p dan q'),
        ('p_i, q_i', 'nilai atribut ke-i dari titik p dan q'),
        ('m', 'jumlah atribut'),
    ])

    h3num(doc, reg, '3.3.2', 'Pembentukan Matriks Jarak')
    body(doc, 'Seluruh jarak antar pasangan data disusun menjadi matriks jarak (proximity matrix) berukuran n x n, bersifat simetris, dan dengan diagonal nol.', indent=False)
    add_equation(doc, 'eq_ag_matriks.png')
    keterangan(doc, [
        ('D', 'matriks jarak'),
        ('n', 'jumlah data'),
        ('D_ij', 'jarak antara data ke-i dan data ke-j'),
    ])

    h3num(doc, reg, '3.3.3', 'Pemilihan Pasangan Klaster Terdekat')
    body(doc, 'Pada setiap iterasi, dipilih pasangan klaster dengan jarak minimum dalam matriks jarak (di luar diagonal).', indent=False)
    add_equation(doc, 'eq_ag_argmin.png')
    keterangan(doc, [
        ('(i*, j*)', 'pasangan indeks klaster terdekat yang akan digabung'),
    ])

    h3num(doc, reg, '3.3.4', 'Metode Linkage')
    body(doc, 'Linkage menentukan cara menghitung jarak antar dua klaster yang masing-masing dapat berisi lebih dari satu anggota. Empat metode linkage yang umum digunakan dijelaskan sebagai berikut.', indent=False)
    body(doc, 'Single linkage memilih jarak minimum antar anggota:')
    add_equation(doc, 'eq_ag_single.png')
    body(doc, 'Complete linkage memilih jarak maksimum antar anggota:')
    add_equation(doc, 'eq_ag_complete.png')
    body(doc, 'Average linkage memilih rata-rata jarak seluruh pasangan anggota:')
    add_equation(doc, 'eq_ag_average.png')
    body(doc, 'Ward linkage meminimalkan kenaikan total jarak kuadrat (SSE) saat menggabungkan dua klaster sehingga klaster cenderung seimbang.')

    h3num(doc, reg, '3.3.5', 'Pembaruan Matriks Jarak (Single Linkage)')
    body(doc, 'Setelah klaster C_i dan C_j digabung menjadi klaster baru, jarak gabungan terhadap klaster lain C_k dihitung sebagai nilai minimum dari kedua jarak asal sesuai aturan single linkage.', indent=False)
    add_equation(doc, 'eq_ag_update.png')
    keterangan(doc, [
        ('C_i ∪ C_j', 'klaster hasil penggabungan'),
        ('C_k', 'klaster lain selain pasangan yang digabung'),
    ])

    h3num(doc, reg, '3.3.6', 'Penentuan Jumlah Klaster dari Gap Terbesar')
    body(doc, 'Karena agglomerative selalu berjalan sampai satu klaster, jumlah klaster final ditentukan dengan memotong dendrogram. Salah satu cara objektif adalah memotong pada lompatan (gap) ketinggian penggabungan terbesar.', indent=False)
    add_equation(doc, 'eq_ag_gap.png')
    keterangan(doc, [
        ('k*', 'jumlah klaster optimal'),
        ('h_i', 'ketinggian penggabungan ke-i (terurut menaik)'),
        ('h_{i+1} - h_i', 'selisih (gap) antar ketinggian penggabungan berurutan'),
    ])

    # ============================================================ BAB IV
    doc.add_page_break()
    bab(doc, reg, 'IV', 'HASIL DAN PEMBAHASAN')
    body(doc, 'Bab IV Hasil dan Pembahasan ini menyajikan hasil penerapan ketiga algoritma data mining pada data studi kasus masing-masing, beserta perhitungan manual dan analisis hasilnya. Pembahasan dibagi menjadi tiga subbab, yaitu pembahasan algoritma Apriori, algoritma K-Means, dan algoritma Agglomerative Clustering.')

    # ----- 4.1 Apriori -----
    h2num(doc, reg, '4.1', 'Algoritma Apriori')
    body(doc, 'Algoritma Apriori diterapkan pada 10 transaksi dengan lima produk (susu, gula, teh, roti, kopi) seperti pada Tabel 4.1. Tujuan analisis adalah menemukan aturan asosiasi antarproduk berdasarkan support, confidence, dan nilai final.')
    caption(doc, reg, 'tbl', 'Tabel 4.1 Data 10 Transaksi Pelanggan')
    rows_trx = [(tid, ', '.join(sorted(it))) for tid, it in transaksi.items()]
    make_table(doc, ['TID', 'Items'], rows_trx)

    h3num(doc, reg, '4.1.1', 'Frequent 1-itemset (F1)')
    body(doc, 'F1 dihitung untuk setiap produk dengan menghitung frekuensi kemunculannya di seluruh transaksi. Hasilnya ditampilkan pada Tabel 4.2.')
    caption(doc, reg, 'tbl', 'Tabel 4.2 Hasil F1 (1-itemset)')
    make_table(doc, ['Itemset', 'Frekuensi', 'Support'],
               [(it, f, f'{s:.2f}') for it, f, s in F1])
    body(doc, 'Dari Tabel 4.2 terlihat bahwa gula memiliki support paling tinggi (0.80), diikuti oleh susu (0.60), teh (0.50), serta roti dan kopi yang sama-sama 0.40.')

    h3num(doc, reg, '4.1.2', 'Frequent 2-itemset (F2)')
    body(doc, 'Dari 5 produk, terdapat C(5,2) = 10 kombinasi 2-itemset. Seluruh kombinasi dihitung supportnya seperti pada Tabel 4.3.')
    caption(doc, reg, 'tbl', 'Tabel 4.3 Hasil F2 (2-itemset)')
    make_table(doc, ['Itemset', 'Frekuensi', 'Support'],
               [(it, f, f'{s:.2f}') for it, f, s in F2])
    body(doc, 'Pasangan dengan support tertinggi adalah teh dan gula (support 0.50), menandakan kedua produk ini sering dibeli bersamaan.')

    h3num(doc, reg, '4.1.3', 'Frequent 3-itemset (F3)')
    body(doc, 'Dari 5 produk, terdapat C(5,3) = 10 kombinasi 3-itemset. Seluruh kombinasi dihitung supportnya seperti pada Tabel 4.4.')
    caption(doc, reg, 'tbl', 'Tabel 4.4 Hasil F3 (3-itemset)')
    make_table(doc, ['Itemset', 'Frekuensi', 'Support'],
               [(it, f, f'{s:.2f}') for it, f, s in F3])

    h3num(doc, reg, '4.1.4', 'Aturan Asosiasi dan Ranking')
    body(doc, 'Dari seluruh kombinasi F2 dan F3, dibentuk aturan A → B beserta support, confidence, dan nilai final. Sepuluh aturan dengan nilai final tertinggi ditampilkan pada Tabel 4.5.')
    caption(doc, reg, 'tbl', 'Tabel 4.5 Sepuluh Aturan Asosiasi Terkuat')
    rows_top = []
    for r in TOP_RULES:
        rows_top.append((f"{r['A']} → {r['B']}", f"{r['sup']:.2f}", f"{r['conf']:.2f}", f"{r['final']:.3f}"))
    make_table(doc, ['Aturan', 'Support', 'Confidence', 'Final'], rows_top)
    body(doc, 'Aturan terkuat adalah teh → gula dengan support 0.50, confidence 1.00, dan nilai final 0.500. Hal ini berarti seluruh transaksi yang memuat teh juga memuat gula, sehingga aturan ini paling layak dijadikan rekomendasi cross-selling. Implementasi kode program Apriori ditampilkan pada Kode Program 4.1.')
    caption(doc, reg, 'code', 'Kode Program 4.1 Implementasi Algoritma Apriori')
    code_box(doc, APRIORI_CODE)

    # ----- 4.2 K-Means -----
    doc.add_page_break()
    h2num(doc, reg, '4.2', 'Algoritma K-Means')
    body(doc, 'Algoritma K-Means diterapkan pada 10 data pelanggan dengan atribut Age (umur) dan Income (penghasilan) seperti pada Tabel 4.6. Tujuan analisis adalah menemukan pengelompokan pelanggan berdasarkan kedekatan karakteristik kedua atribut tersebut.')
    caption(doc, reg, 'tbl', 'Tabel 4.6 Data 10 Pelanggan')
    rows_cust = [(int(r['CustID']), int(r['Age']), int(r['Income'])) for _, r in cust.iterrows()]
    make_table(doc, ['CustID', 'Age', 'Income'], rows_cust)
    add_image(doc, 'km_scatter.png', 11)
    caption(doc, reg, 'fig', 'Gambar 4.1 Sebaran 10 Pelanggan')

    h3num(doc, reg, '4.2.1', 'Perhitungan Manual K=2')
    body(doc, 'Untuk K=2, dipilih centroid awal C1 = P1 = (41, 19) dan C2 = P2 = (47, 100). Setiap data dihitung jaraknya terhadap kedua centroid, lalu ditetapkan ke klaster dengan centroid terdekat. Hasil perhitungan jarak iterasi pertama ditampilkan pada Tabel 4.7.')
    caption(doc, reg, 'tbl', 'Tabel 4.7 Jarak Iterasi Pertama K-Means (K=2)')
    make_table(doc, ['CustID', 'Age', 'Income', 'Jarak ke C1', 'Jarak ke C2', 'Klaster'],
               [(c[0], c[1], c[2], f'{c[3]:.2f}', f'{c[4]:.2f}', c[5]) for c in km_iter1])
    body(doc, 'Setelah penetapan klaster, centroid diperbarui dengan menghitung rata-rata koordinat anggotanya. Proses ini diulang hingga konvergen. Ringkasan posisi centroid setiap iterasi ditampilkan pada Tabel 4.8.')
    caption(doc, reg, 'tbl', 'Tabel 4.8 Ringkasan Iterasi K-Means (K=2)')
    rows_iter = []
    for h in km_hist2:
        c1 = h['cent'][0]; c2 = h['cent'][1]
        rows_iter.append((h['iter'], f'({c1[0]:.2f}, {c1[1]:.2f})', f'({c2[0]:.2f}, {c2[1]:.2f})',
                          ' '.join(map(str, h['labels']))))
    make_table(doc, ['Iterasi', 'Centroid C1', 'Centroid C2', 'Label'], rows_iter)
    body(doc, f'Algoritma konvergen setelah {len(km_hist2)} iterasi dengan centroid final C1 = ({km_cent2[0][0]:.2f}, {km_cent2[0][1]:.2f}) dan C2 = ({km_cent2[1][0]:.2f}, {km_cent2[1][1]:.2f}). Visualisasi hasil ditampilkan pada Gambar 4.2.')
    add_image(doc, 'km_k2.png', 11)
    caption(doc, reg, 'fig', 'Gambar 4.2 Hasil K-Means Manual untuk K=2')

    h3num(doc, reg, '4.2.2', 'Penentuan K Optimal dengan Elbow Method')
    body(doc, 'Untuk menentukan nilai K yang paling sesuai, dihitung SSE untuk K=1 hingga K=7 pada data yang telah distandarisasi (Z-Score) seperti pada Tabel 4.9. Kurva SSE diplot pada Gambar 4.3.')
    caption(doc, reg, 'tbl', 'Tabel 4.9 Nilai SSE per K (Elbow Method)')
    make_table(doc, ['K', 'SSE (data ter-standarisasi)'], [(k, f'{s:.4f}') for k, s in KM_SSE])
    add_image(doc, 'km_elbow.png', 12)
    caption(doc, reg, 'fig', 'Gambar 4.3 Elbow Method untuk Penentuan K Optimal')
    body(doc, 'Dari Gambar 4.3 terlihat bahwa penurunan SSE paling tajam terjadi dari K=1 ke K=2, lalu mulai melandai pada K=3. Titik siku (elbow) berada di sekitar K=2 atau K=3, sehingga kedua nilai ini menjadi kandidat K optimal. Hasil clustering untuk K=3 hingga K=6 ditampilkan pada Gambar 4.4.')
    add_image(doc, 'km_k36.png', 13)
    caption(doc, reg, 'fig', 'Gambar 4.4 Hasil K-Means untuk K=3 sampai K=6')

    h3num(doc, reg, '4.2.3', 'Kode Program')
    body(doc, 'Implementasi K-Means manual dan otomatis ditampilkan pada Kode Program 4.2.')
    caption(doc, reg, 'code', 'Kode Program 4.2 Implementasi K-Means dengan Elbow Method')
    code_box(doc, KMEANS_CODE)

    # ----- 4.3 Agglomerative -----
    doc.add_page_break()
    h2num(doc, reg, '4.3', 'Algoritma Agglomerative Clustering')
    body(doc, 'Algoritma Agglomerative Hierarchical Clustering diterapkan pada 10 data mahasiswa dengan atribut Nilai Tugas dan Nilai Ujian seperti pada Tabel 4.10. Tujuan analisis adalah menemukan pengelompokan mahasiswa secara hirarkis dan menentukan jumlah klaster dari dendrogram.')
    caption(doc, reg, 'tbl', 'Tabel 4.10 Data 10 Mahasiswa')
    rows_agg = [(r['Mhs'], int(r['Tugas']), int(r['Ujian'])) for _, r in data.iterrows()]
    make_table(doc, ['Mahasiswa', 'Nilai Tugas', 'Nilai Ujian'], rows_agg)
    add_image(doc, 'ag_scatter.png', 11)
    caption(doc, reg, 'fig', 'Gambar 4.5 Sebaran 10 Mahasiswa')

    h3num(doc, reg, '4.3.1', 'Pembentukan Matriks Jarak Awal')
    body(doc, 'Langkah pertama adalah menghitung jarak Euclidean antar seluruh pasangan data. Karena terdapat 10 data, matriks jarak berukuran 10 x 10 dan bersifat simetris dengan diagonal nol. Matriks jarak awal ditampilkan pada Tabel 4.11.')
    caption(doc, reg, 'tbl', 'Tabel 4.11 Matriks Jarak Awal Antar Mahasiswa (Euclidean)')
    t = doc.add_table(rows=1, cols=11); t.alignment = WD_TABLE_ALIGNMENT.CENTER; set_table_borders(t)
    cell(t.rows[0].cells[0], '', bold=True, size=9); set_cell_bg(t.rows[0].cells[0], 'D9D9D9')
    for j, nm in enumerate(names, 1):
        cell(t.rows[0].cells[j], nm, bold=True, size=9); set_cell_bg(t.rows[0].cells[j], 'D9D9D9')
    for i, nm in enumerate(names):
        cs = t.add_row().cells
        cell(cs[0], nm, bold=True, size=9); set_cell_bg(cs[0], 'D9D9D9')
        for j in range(10):
            cell(cs[j + 1], f'{dist_mat[i, j]:.2f}', size=9)

    h3num(doc, reg, '4.3.2', 'Perhitungan Manual Single Linkage')
    body(doc, 'Pada setiap iterasi, dicari pasangan dengan jarak minimum di matriks, kedua klaster digabung, lalu matriks diperbarui menggunakan rumus single linkage (nilai minimum dari kedua baris/kolom yang digabung). Proses diulang sembilan kali hingga seluruh data tergabung. Ringkasan langkah penggabungan ditampilkan pada Tabel 4.12.')
    caption(doc, reg, 'tbl', 'Tabel 4.12 Ringkasan Sembilan Iterasi Penggabungan (Single Linkage)')
    make_table(doc, ['Iterasi', 'Klaster A', 'Klaster B', 'Jarak Merge', 'Sisa Klaster'],
               [(it, a, b, f'{d:.3f}', sisa) for it, a, b, d, sisa in merge_rows])
    body(doc, 'Pada Tabel 4.12 terlihat bahwa enam penggabungan pertama terjadi pada jarak 1.000 karena masih menggabungkan data dalam kelompok yang sama. Penggabungan ketujuh terjadi pada jarak 1.414, sedangkan penggabungan kedelapan melonjak ke 3.606. Lonjakan inilah yang menjadi dasar penentuan jumlah klaster pada subbab berikutnya.')

    h3num(doc, reg, '4.3.3', 'Kode Program')
    body(doc, 'Implementasi Agglomerative Clustering manual dan verifikasi dengan SciPy ditampilkan pada Kode Program 4.3.')
    caption(doc, reg, 'code', 'Kode Program 4.3 Implementasi Agglomerative Clustering Manual')
    code_box(doc, AGG_CODE)

    h3num(doc, reg, '4.3.4', 'Hasil dan Penentuan Jumlah Klaster')
    body(doc, 'Hasil seluruh proses penggabungan divisualisasikan dalam bentuk dendrogram pada Gambar 4.6. Garis putus-putus merah menunjukkan titik potong otomatis pada ketinggian sekitar 2.51.')
    add_image(doc, 'ag_dendrogram.png', 13)
    caption(doc, reg, 'fig', 'Gambar 4.6 Dendrogram Agglomerative Single Linkage')
    body(doc, 'Penentuan jumlah klaster dilakukan dengan mengamati lompatan ketinggian penggabungan. Urutan ketinggian adalah 1.0 (enam kali), 1.414, 3.606, dan 4.243. Lompatan terbesar bernilai 2.191 yaitu antara 1.414 dan 3.606. Lompatan ini terjadi ketika algoritma mulai menggabungkan dua kelompok berbeda, sehingga titik potong diletakkan tepat di bawahnya dan menghasilkan tiga klaster, bukan dua. Hal ini menegaskan bahwa algoritma Agglomerative tidak berhenti di dua klaster, melainkan berjalan sampai satu klaster, dan jumlah klaster final ditentukan oleh karakteristik data. Hasil pengelompokan tiga klaster ditampilkan pada Gambar 4.7.')
    add_image(doc, 'ag_hasil.png', 11)
    caption(doc, reg, 'fig', 'Gambar 4.7 Hasil Pengelompokan Tiga Klaster')
    body(doc, 'Untuk menegaskan bahwa jumlah klaster bergantung pada letak pemotongan dendrogram, dilakukan pemotongan menjadi 2, 3, dan 4 klaster. Hasilnya ditampilkan pada Gambar 4.8.')
    add_image(doc, 'ag_potong.png', 14)
    caption(doc, reg, 'fig', 'Gambar 4.8 Perbandingan Pemotongan Dendrogram menjadi 2, 3, dan 4 Klaster')
    body(doc, 'Pada Gambar 4.8 terlihat bahwa pemotongan menjadi 2 klaster menggabungkan kelompok menengah (D, E, F) ke salah satu kelompok besar sehingga kurang merepresentasikan struktur data yang sebenarnya. Pemotongan menjadi 3 klaster menghasilkan pemisahan yang paling bersih, di mana ketiga kelompok yaitu nilai rendah {A, B, C}, nilai menengah {D, E, F}, dan nilai tinggi {G, H, I, J} terpisah dengan jelas sesuai sebaran data asli. Pemotongan menjadi 4 klaster mulai memecah salah satu kelompok sehingga terjadi over-segmentasi, di mana salah satu klaster hanya berisi sedikit anggota. Hal ini menegaskan bahwa pemilihan letak pemotongan dendrogram secara langsung menentukan jumlah dan kualitas klaster final, dan untuk data ini tiga klaster merupakan pilihan paling sesuai.')

    body(doc, 'Selain pengaruh letak pemotongan, dilakukan pula perbandingan empat metode linkage (single, complete, average, dan ward) untuk melihat apakah pilihan metode mempengaruhi hasil klaster. Hasilnya ditampilkan pada Gambar 4.9.')
    add_image(doc, 'ag_linkage.png', 13)
    caption(doc, reg, 'fig', 'Gambar 4.9 Perbandingan Empat Metode Linkage')
    body(doc, 'Pada Gambar 4.9 terlihat bahwa keempat metode linkage menghasilkan struktur hirarki yang serupa, yaitu tiga klaster yang sama: {A, B, C}, {D, E, F}, dan {G, H, I, J}. Perbedaan utama terletak pada skala ketinggian penggabungan. Metode single linkage menghasilkan jarak penggabungan paling rendah karena mengambil nilai minimum antar anggota, sedangkan complete linkage menghasilkan skala paling tinggi karena memilih jarak maksimum antar anggota. Metode average dan ward berada di antaranya dengan pola yang lebih halus. Karena ketiga kelompok pada data terpisah cukup jelas, pilihan metode linkage tidak mengubah hasil klaster final, melainkan hanya mempengaruhi skala ketinggian pada dendrogram. Hal ini menunjukkan bahwa pada data dengan kelompok yang terpisah baik, Agglomerative Clustering bersifat robust terhadap pilihan metode linkage.')

    # ============================================================ BAB V
    doc.add_page_break()
    bab(doc, reg, 'V', 'PENUTUP')

    h2num(doc, reg, '5.1', 'Kesimpulan')
    body(doc, 'Berdasarkan hasil dan pembahasan implementasi ketiga algoritma data mining pada laporan ini, dapat ditarik beberapa kesimpulan sebagai berikut.', indent=False)
    numbered(doc, 1, 'Algoritma Apriori berhasil menemukan aturan asosiasi pada 10 transaksi dengan lima produk. Aturan terkuat adalah teh → gula dengan support 0.50, confidence 1.00, dan nilai final 0.500. Sepuluh aturan teratas dapat dimanfaatkan sebagai dasar rekomendasi cross-selling.')
    numbered(doc, 2, 'Algoritma K-Means konvergen dalam empat iterasi untuk K=2 dengan centroid awal P1 dan P2. Elbow Method menunjukkan titik siku berada di sekitar K=2 atau K=3, sehingga kedua nilai ini menjadi kandidat K optimal sesuai karakteristik data. Standarisasi atribut (Z-Score) penting karena skala Income jauh lebih besar daripada Age.')
    numbered(doc, 3, 'Algoritma Agglomerative Hierarchical Clustering berhasil membentuk hirarki klaster pada 10 mahasiswa dengan sembilan iterasi penggabungan menggunakan single linkage. Matriks jarak diperbarui pada setiap iterasi dengan rumus minimum. Jumlah klaster final ditentukan dari lompatan ketinggian penggabungan terbesar, yang menghasilkan tiga klaster, bukan dua. Hasil ini sekaligus menegaskan bahwa jumlah klaster pada Agglomerative bergantung pada karakteristik data dan letak pemotongan dendrogram.')
    numbered(doc, 4, 'Ketiga algoritma memiliki tujuan berbeda. Apriori untuk menemukan keterkaitan antaritem, K-Means untuk pengelompokan dengan jumlah klaster yang ditentukan di awal, sedangkan Agglomerative Clustering untuk membentuk hirarki klaster lengkap tanpa perlu menentukan jumlah klaster terlebih dahulu. Pemilihan algoritma yang tepat bergantung pada karakteristik data dan tujuan analisis.')

    h2num(doc, reg, '5.2', 'Saran')
    body(doc, 'Beberapa saran untuk pengembangan penelitian selanjutnya adalah sebagai berikut.', indent=False)
    numbered(doc, 1, 'Penelitian dapat dilakukan pada dataset yang lebih besar dan kompleks untuk melihat kinerja ketiga algoritma pada data nyata, sekaligus mengevaluasi keterbatasan komputasinya.')
    numbered(doc, 2, 'Algoritma Apriori dapat dikembangkan dengan menambahkan ambang batas minimum support dan minimum confidence, serta dibandingkan dengan algoritma asosiasi lain seperti FP-Growth.')
    numbered(doc, 3, 'Algoritma K-Means dapat dievaluasi dengan ukuran kualitas klaster seperti Silhouette Score, serta dibandingkan dengan algoritma clustering lain seperti K-Medoids dan DBSCAN.')
    numbered(doc, 4, 'Algoritma Agglomerative Clustering dapat dibandingkan dengan berbagai metode linkage secara lebih mendalam, dan penentuan jumlah klaster dapat dikombinasikan dengan ukuran evaluasi seperti Silhouette Score atau Davies-Bouldin Index agar lebih objektif.')

    # ============================================================ DAFTAR PUSTAKA
    doc.add_page_break()
    special_heading(doc, reg, 'DAFTAR PUSTAKA')
    for r in [
        'Agrawal, R., & Srikant, R. (1994). Fast Algorithms for Mining Association Rules in Large Databases. Proceedings of the 20th International Conference on Very Large Data Bases (VLDB), 487-499.',
        'Han, J., Kamber, M., & Pei, J. (2012). Data Mining: Concepts and Techniques (3rd ed.). Waltham: Morgan Kaufmann Publishers.',
        'Hartigan, J. A., & Wong, M. A. (1979). A K-Means Clustering Algorithm. Journal of the Royal Statistical Society. Series C (Applied Statistics), 28(1), 100-108.',
        'Larose, D. T., & Larose, C. D. (2014). Discovering Knowledge in Data: An Introduction to Data Mining (2nd ed.). New Jersey: John Wiley & Sons Inc.',
        'MacQueen, J. (1967). Some Methods for Classification and Analysis of Multivariate Observations. Proceedings of the Fifth Berkeley Symposium on Mathematical Statistics and Probability, 1, 281-297.',
        'Tan, P.-N., Steinbach, M., Karpatne, A., & Kumar, V. (2019). Introduction to Data Mining (2nd ed.). New York: Pearson.',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(r)


# ============================================================ FRONT MATTER
def build_cover(doc):
    plain_center(doc, 'LAPORAN UJIAN TENGAH SEMESTER', bold=True, sb=24)
    plain_center(doc, 'MATA KULIAH DATA MINING', bold=True, sa=18)
    pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(os.path.join(FIG, 'logo_udayana.png'), width=Cm(4.5))
    plain_center(doc, '', sa=12)
    plain_center(doc, 'DOSEN PENGAMPU', bold=True, sb=12)
    plain_center(doc, 'Dr. Eng. I Putu Agung Bayupati, S.T., M.T.', sa=18)
    plain_center(doc, 'DISUSUN OLEH:', bold=True, sb=12)
    plain_center(doc, 'KELOMPOK 10', bold=True)
    for nm, nim in [('Ravi Arnan Irianto', '2305551076'),
                    ('Ezza Putra Wibawa', '2305551144'),
                    ('Ketut Riski Prananda', '2005551125')]:
        plain_center(doc, f'{nm}\t\t({nim})')
    plain_center(doc, '', sa=18)
    plain_center(doc, 'PROGRAM STUDI TEKNOLOGI INFORMASI', bold=True, sb=18)
    plain_center(doc, 'FAKULTAS TEKNIK', bold=True)
    plain_center(doc, 'UNIVERSITAS UDAYANA', bold=True)
    plain_center(doc, 'BALI', bold=True)
    plain_center(doc, '2026', bold=True)


def build_front(doc, reg, pages):
    plain_center(doc, 'DAFTAR ISI', bold=True, sa=12)
    for e in reg:
        if e['cat'] == 'toc':
            toc_line(doc, e['label'], pages.get(e['token'], 0), e['level'])
    doc.add_page_break()
    plain_center(doc, 'DAFTAR GAMBAR', bold=True, sa=12)
    for e in reg:
        if e['cat'] == 'fig':
            toc_line(doc, e['label'], pages.get(e['token'], 0), 0)
    doc.add_page_break()
    plain_center(doc, 'DAFTAR KODE PROGRAM', bold=True, sa=12)
    for e in reg:
        if e['cat'] == 'code':
            toc_line(doc, e['label'], pages.get(e['token'], 0), 0)
    doc.add_page_break()
    plain_center(doc, 'DAFTAR TABEL', bold=True, sa=12)
    for e in reg:
        if e['cat'] == 'tbl':
            toc_line(doc, e['label'], pages.get(e['token'], 0), 0)


# ============================================================ ASSEMBLE
def assemble(pages):
    doc = Document()
    set_base_style(doc)
    style_heading(doc, 'Heading 1', 14)
    style_heading(doc, 'Heading 2', 12)
    style_heading(doc, 'Heading 3', 12)
    s = doc.sections[0]
    s.top_margin = Cm(3); s.bottom_margin = Cm(3)
    s.left_margin = Cm(4); s.right_margin = Cm(3)
    s.page_height = Cm(29.7); s.page_width = Cm(21)
    build_cover(doc)
    s.footer.is_linked_to_previous = False

    doc.add_section(WD_SECTION.NEW_PAGE)
    s2 = doc.sections[1]
    s2.top_margin = Cm(3); s2.bottom_margin = Cm(3); s2.left_margin = Cm(4); s2.right_margin = Cm(3)
    reg2 = []; build_body(Document(), reg2)
    build_front(doc, reg2, pages)
    page_footer(s2, 'lowerRoman', 1)

    doc.add_section(WD_SECTION.NEW_PAGE)
    s3 = doc.sections[2]
    s3.top_margin = Cm(3); s3.bottom_margin = Cm(3); s3.left_margin = Cm(4); s3.right_margin = Cm(3)
    page_footer(s3, 'decimal', 1)
    reg3 = []; build_body(doc, reg3)
    return doc


if __name__ == '__main__':
    mode = sys.argv[1] if len(sys.argv) > 1 else 'passB'
    if mode == 'passA':
        doc = assemble(pages={})
        doc.save(TMP_DOCX)
        print('passA saved:', TMP_DOCX)
    else:
        reg = []
        build_body(Document(), reg)
        pages = detect_pages(TMP_PDF, reg)
        doc = assemble(pages)
        doc.save(OUT_DOCX)
        print('Final saved:', OUT_DOCX)
        sample = list(reg)[:6]
        print('Detected pages sample:', {e['label']: pages.get(e['token']) for e in sample})
