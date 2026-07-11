"""Build ONLY subbab 4.4 (analisis centroid awal di titik terjauh) as a docx,
matching the clean report style (A4, margin 4/3/3/3 cm, Times New Roman 12,
justified, 1.5 spacing). No footer page numbers (added later by overlay)."""
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = 'Times New Roman'
FIG = '/home/ravi/Projects/DataMining/docs/figures/km_terjauh.png'

# ---- hitung skenario (sama seperti notebook/laporan v2) ----
Age = [41, 47, 33, 29, 47, 40, 38, 42, 26, 47]
Income = [19, 100, 57, 19, 253, 81, 56, 64, 18, 115]
X = np.array(list(zip(Age, Income)), dtype=float)
namesP = [f'P{i+1}' for i in range(len(X))]

def euclid(a, b):
    return np.sqrt(np.sum((np.array(a) - np.array(b)) ** 2))

def kmeans(Xd, init, mx=20):
    c = np.array(init, float); hist = 0
    for _ in range(mx):
        lab = np.array([int(np.argmin([euclid(x, ci) for ci in c])) for x in Xd])
        hist += 1
        nc = np.array([Xd[lab == k].mean(axis=0) if (lab == k).any() else c[k] for k in range(len(c))])
        if np.allclose(nc, c):
            break
        c = nc
    return c, lab, hist

mean = X.mean(axis=0)
jd = np.sqrt(((X - mean) ** 2).sum(axis=1))
i_far, i_cen = int(jd.argmax()), int(jd.argmin())
C2 = X[1]
cD, lD, hD = kmeans(X, [X[i_cen], C2])
cL, lL, hL = kmeans(X, [X[i_far], C2])
def mem(lab, k): return ', '.join(namesP[i] for i in range(len(lab)) if lab[i] == k)
def fmt(p): return f'({int(round(p[0]))}, {int(round(p[1]))})'
rD = f'{int((lD==0).sum())} : {int((lD==1).sum())}'
rL = f'{int((lL==0).sum())} : {int((lL==1).sum())}'

# ---- dokumen ----
doc = Document()
st = doc.styles['Normal']
st.font.name = FONT; st.font.size = Pt(12); st.font.color.rgb = RGBColor(0, 0, 0)
rf = st.element.get_or_add_rPr().get_or_add_rFonts()
for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
    rf.set(qn(a), FONT)
st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
st.paragraph_format.space_after = Pt(0)

sec = doc.sections[0]
sec.page_height = Cm(29.7); sec.page_width = Cm(21)
sec.top_margin = Cm(3); sec.bottom_margin = Cm(3)
sec.left_margin = Cm(4); sec.right_margin = Cm(3)

def run(p, text, bold=False, italic=False, size=12):
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.name = FONT; r.font.size = Pt(size)
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:cs'), FONT)
    return r

def heading():
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(1.1), WD_TAB_ALIGNMENT.LEFT)
    run(p, '4.4', bold=True); run(p, '\tAnalisis Pengaruh Centroid Awal terhadap Outlier', bold=True)

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.0)
    run(p, text)

def numbered(n, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    run(p, f'{n}.\t{text}')

def caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    run(p, text, bold=True, size=11)

def set_cell_bg(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), color)
    tcPr.append(sh)

def set_borders(tbl):
    tblPr = tbl._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4'); e.set(qn('w:space'), '0'); e.set(qn('w:color'), '000000')
        b.append(e)
    tblPr.append(b)

def cell_text(c, text, bold=False, size=10.5, center=True):
    c.text = ''
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(0)
    run(p, str(text), bold=bold, size=size)

# ---- isi ----
heading()
body(f'Perhitungan pada subbab 4.2.1 memakai centroid awal yang berada di dalam sebaran '
     f'data. Untuk menguji sensitivitas K-Means terhadap pemilihan centroid awal, dilakukan '
     f'percobaan tambahan dengan menempatkan salah satu centroid awal tepat di titik terjauh '
     f'(outlier) dari kerumunan data. Jarak setiap titik terhadap rata-rata data dihitung, dan '
     f'diperoleh {namesP[i_far]} = {fmt(X[i_far])} sebagai titik terjauh (berada di luar '
     f'layer atau kerumunan) serta {namesP[i_cen]} = {fmt(X[i_cen])} sebagai titik paling '
     f'sentral (di dalam layer). Dua skenario dibandingkan, dengan centroid kedua sengaja '
     f'dibuat sama (C2 tetap = P2) agar perbedaan hasil murni disebabkan oleh perpindahan '
     f'centroid pertama.')
numbered(1, f'Skenario A (di dalam layer): centroid awal C1 = {namesP[i_cen]} {fmt(X[i_cen])} '
            f'dan C2 = P2 {fmt(C2)}.')
numbered(2, f'Skenario B (di luar layer): centroid awal C1 = {namesP[i_far]} {fmt(X[i_far])} '
            f'(titik terjauh) dan C2 = P2 {fmt(C2)}.')
body('Percobaan ini menjawab dua pertanyaan, yaitu apa yang berbeda jika centroid awal '
     'diletakkan di titik terjauh atau outlier, dan apakah anggota klaster berubah atau tetap '
     'sama. Perbandingan kedua skenario ditampilkan pada Tabel 4.13 dan divisualisasikan pada '
     'Gambar 4.10.')

caption('Tabel 4.13 Perbandingan Centroid Awal di Dalam vs di Luar Layer')
rows = [
    ('Aspek', 'A – Dalam Layer', 'B – Luar Layer'),
    ('Centroid awal C1', f'{namesP[i_cen]} {fmt(X[i_cen])}', f'{namesP[i_far]} {fmt(X[i_far])}'),
    ('Centroid awal C2 (tetap)', f'P2 {fmt(C2)}', f'P2 {fmt(C2)}'),
    ('Posisi C1', 'di dalam kerumunan', 'di outlier (luar layer)'),
    ('Jumlah iterasi', str(hD), str(hL)),
    ('Centroid akhir C1', fmt(cD[0]), fmt(cL[0])),
    ('Centroid akhir C2', fmt(cD[1]), fmt(cL[1])),
    ('Rasio anggota (K1 : K2)', rD, rL),
    ('Anggota Klaster 1', mem(lD, 0), mem(lL, 0)),
    ('Anggota Klaster 2', mem(lD, 1), mem(lL, 1)),
]
tbl = doc.add_table(rows=len(rows), cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_borders(tbl)
tbl.columns[0].width = Cm(4.2); tbl.columns[1].width = Cm(5.6); tbl.columns[2].width = Cm(5.6)
for ri, r in enumerate(rows):
    for ci, v in enumerate(r):
        c = tbl.rows[ri].cells[ci]
        left = (ci == 0)
        cell_text(c, v, bold=(ri == 0), center=not left if ri != 0 else True)
        if ri == 0:
            set_cell_bg(c, 'D9D9D9')

# gambar
pimg = doc.add_paragraph(); pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
pimg.paragraph_format.space_before = Pt(6)
pimg.add_run().add_picture(FIG, width=Cm(14))
caption('Gambar 4.10 Perbandingan Centroid Awal di Dalam vs di Luar Layer')

body(f'Jawaban pertanyaan pertama, yaitu ketika centroid awal diletakkan di titik terjauh '
     f'{namesP[i_far]}, titik tersebut langsung menjadi klaster yang berisi satu anggota saja. '
     f'Pembagian yang semula wajar ({rD}, memisahkan kelompok Income rendah dan Income tinggi) '
     f'berubah menjadi timpang ({rL}), yaitu satu klaster hanya berisi outlier dan satu klaster '
     f'menampung sembilan titik sisanya. Hal ini terjadi karena {namesP[i_far]} memiliki Income '
     f'yang jauh lebih tinggi (253) dibanding titik lain (18 sampai 115), sehingga tidak ada '
     f'titik lain yang lebih dekat kepadanya. Karena anggota klaster tersebut hanya satu titik, '
     f'pembaruan centroid (rata-rata anggota) mengembalikan posisi yang sama, sehingga centroid '
     f'terkunci di outlier.')
body(f'Jawaban pertanyaan kedua, yaitu anggota klaster berubah total. Dari rasio {rD} pada '
     f'skenario di dalam layer menjadi {rL} pada skenario di luar layer, padahal yang diubah '
     f'hanya satu centroid awal (C2 dibuat tetap). Hal ini membuktikan bahwa K-Means sangat '
     f'sensitif terhadap posisi centroid awal dan keberadaan outlier. Menaruh centroid awal di '
     f'titik terjauh menghasilkan klaster yang tidak informatif. Oleh karena itu, praktik yang '
     f'dianjurkan adalah melakukan standarisasi data agar skala Income tidak mendominasi, '
     f'menggunakan inisialisasi k-means++ agar centroid awal tersebar dan menghindari outlier, '
     f'serta mengulang inisialisasi beberapa kali (n_init) dan memilih hasil terbaik.')

out = '/home/ravi/Projects/DataMining/docs/_sec44.docx'
doc.save(out)
print('saved', out, '| rasio', rD, '->', rL, '| iter', hD, hL)
