"""Dokumen Word penjelasan ALUR PER CELL — file TERPISAH per tugas.

Mengonversi tiap notebook menjadi dokumen Word yang menampilkan setiap cell
kode beserta penjelasannya (mengikuti urutan & isi notebook).

Menghasilkan:
  - Alur_Tugas1_Apriori.docx
  - Alur_Tugas2_KMeans.docx
  - Alur_Tugas3_Agglomerative.docx

Pakai:  .venv/bin/python docs/build_alur.py
"""
import os
import re
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
FONT = 'Times New Roman'
MONO = 'Courier New'
ACCENT = RGBColor(0x1F, 0x39, 0x64)
CODE_BG = 'F2F3F5'

JOBS = [
    ('Tugas1_Apriori.ipynb', 'Alur_Tugas1_Apriori.docx',
     'Algoritma Apriori (Aturan Asosiasi)'),
    ('Tugas2_KMeans_Elbow.ipynb', 'Alur_Tugas2_KMeans.docx',
     'K-Means Clustering + Elbow Method'),
    ('Tugas3_Agglomerative.ipynb', 'Alur_Tugas3_Agglomerative.docx',
     'Agglomerative Hierarchical Clustering'),
]


# ---------------------------------------------------------------- font helper
def set_font(run, size=12, bold=False, italic=False, color=None, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font)


def shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


# ---------------------------------------------------------------- math cleanup
def clean_math(s):
    s = s.replace('\\text', '').replace('\\mathrm', '')
    s = re.sub(r'\\dfrac\{([^{}]*)\}\{([^{}]*)\}', r'(\1)/(\2)', s)
    s = re.sub(r'\\frac\{([^{}]*)\}\{([^{}]*)\}', r'(\1)/(\2)', s)
    s = re.sub(r'\\binom\{([^{}]*)\}\{([^{}]*)\}', r'C(\1,\2)', s)
    s = (s.replace('\\times', ' x ').replace('\\cup', ' U ')
           .replace('\\cap', ' ∩ ').replace('\\min', 'min').replace('\\max', 'max')
           .replace('\\sqrt', 'akar').replace('\\sum', 'Σ').replace('\\in', ' ∈ ')
           .replace('\\le', '≤').replace('\\ge', '≥').replace('\\cdot', '·')
           .replace('\\approx', '≈').replace('\\rightarrow', '->').replace('\\to', '->'))
    s = re.sub(r'\\[a-zA-Z]+', '', s)       # drop remaining commands
    s = s.replace('{', '').replace('}', '').replace('$', '')
    return s


# ---------------------------------------------------------------- inline runs
INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+`|\$[^$]+\$)')


def add_inline(p, text):
    text = text.replace('\\(', '$').replace('\\)', '$')
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            set_font(p.add_run(tok[2:-2]), bold=True)
        elif tok.startswith('`') and tok.endswith('`'):
            set_font(p.add_run(tok[1:-1]), size=11, font=MONO)
        elif tok.startswith('$') and tok.endswith('$'):
            set_font(p.add_run(clean_math(tok)), italic=True)
        else:
            set_font(p.add_run(clean_math(tok)))


# ---------------------------------------------------------------- block render
def render_markdown(doc, src):
    lines = src.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        # tables
        if line.lstrip().startswith('|') and i + 1 < len(lines) \
                and set(lines[i + 1].replace('|', '').strip()) <= set('-: '):
            tbl_lines = []
            while i < len(lines) and lines[i].lstrip().startswith('|'):
                tbl_lines.append(lines[i])
                i += 1
            render_table(doc, tbl_lines)
            continue

        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text = clean_math(re.sub(r'[*`]', '', m.group(2)))
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
            if level <= 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_font(p.add_run(text), size=17, bold=True, color=ACCENT)
            elif level == 2:
                set_font(p.add_run(text), size=14, bold=True, color=ACCENT)
            else:
                set_font(p.add_run(text), size=12, bold=True)
            i += 1
            continue

        mb = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if mb:
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, mb.group(2))
            i += 1
            continue

        mn = re.match(r'^\s*\d+\.\s+(.*)', line)
        if mn:
            p = doc.add_paragraph(style='List Number')
            add_inline(p, mn.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1


def render_table(doc, tbl_lines):
    rows = []
    for ln in tbl_lines:
        cells = [c.strip() for c in ln.strip().strip('|').split('|')]
        rows.append(cells)
    if len(rows) < 2:
        return
    header = rows[0]
    body = rows[2:]
    t = doc.add_table(rows=1, cols=len(header))
    t.style = 'Light Grid Accent 1'
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]
        c.text = ''
        set_font(c.paragraphs[0].add_run(clean_math(re.sub(r'[*`]', '', h))),
                 size=11, bold=True)
    for r in body:
        cells = t.add_row().cells
        for j in range(len(header)):
            val = r[j] if j < len(r) else ''
            cells[j].text = ''
            set_font(cells[j].paragraphs[0].add_run(clean_math(re.sub(r'[*`]', '', val))),
                     size=11)
    doc.add_paragraph()


def render_code(doc, src):
    code = src.rstrip('\n')
    if not code.strip():
        return
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    set_font(cap.add_run('Kode:'), size=11, bold=True, color=ACCENT)
    for ln in code.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        shade(p, CODE_BG)
        set_font(p.add_run(ln if ln else ' '), size=10, font=MONO)
    doc.add_paragraph()


# ---------------------------------------------------------------- doc builder
def new_doc():
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    return doc


def build(nb_name, out_name, subtitle):
    with open(os.path.join(ROOT, nb_name), encoding='utf-8') as f:
        nb = json.load(f)
    doc = new_doc()

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(sub.add_run('Penjelasan Alur Per Cell — ' + subtitle), size=13, bold=True)
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(sub2.add_run('Data Mining — Kelompok 10'), size=12)
    doc.add_paragraph()

    for cell in nb['cells']:
        src = ''.join(cell['source'])
        if cell['cell_type'] == 'markdown':
            render_markdown(doc, src)
        elif cell['cell_type'] == 'code':
            render_code(doc, src)

    out = os.path.join(HERE, out_name)
    doc.save(out)
    return out


if __name__ == '__main__':
    for nb_name, out_name, subtitle in JOBS:
        print('Tersimpan:', build(nb_name, out_name, subtitle))
